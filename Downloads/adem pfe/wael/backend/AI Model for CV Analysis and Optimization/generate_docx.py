from docx import Document

def create_sample_cv():
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Software Engineer')
    doc.add_paragraph('Email: john.doe@example.com | LinkedIn: linkedin.com/in/johndoe | GitHub: github.com/johndoe')

    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Highly motivated Software Engineer with 5+ years of experience in building scalable web applications. Proficient in Python, Java, and cloud technologies.')

    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Senior Software Engineer | Tech Solutions Inc. | 2020 - Present')
    doc.add_paragraph('- Led a team of 5 developers to deliver a high-traffic e-commerce platform.', style='List Bullet')
    doc.add_paragraph('- Increased system performance by 30% through database optimization and caching strategies.', style='List Bullet')
    doc.add_paragraph('- Developed and maintained RESTful APIs using FastAPI and PostgreSQL.', style='List Bullet')
    doc.add_paragraph('- Implemented CI/CD pipelines using GitHub Actions and Docker.', style='List Bullet')

    doc.add_paragraph('Software Engineer | Innovative Apps Co. | 2018 - 2020')
    doc.add_paragraph('- Built a real-time data visualization dashboard using React and Node.js.', style='List Bullet')
    doc.add_paragraph('- Reduced server response time by 40% by optimizing backend logic.', style='List Bullet')
    doc.add_paragraph('- Collaborated with cross-functional teams to define project requirements and timelines.', style='List Bullet')

    doc.add_heading('Education', level=1)
    doc.add_paragraph('Bachelor of Science in Computer Science | University of Technology | 2014 - 2018')

    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Languages: Python, Java, JavaScript, SQL, HTML/CSS')
    doc.add_paragraph('Frameworks: FastAPI, React, Node.js, Express')
    doc.add_paragraph('Tools: Docker, Kubernetes, Git, AWS, GitHub Actions')
    doc.add_paragraph('Soft Skills: Leadership, Teamwork, Problem Solving, Communication')

    doc.add_heading('Projects', level=1)
    doc.add_paragraph('Personal Portfolio Website: Built using React and hosted on AWS S3.')
    doc.add_paragraph('Open Source Contributor: Contributed to several popular Python libraries on GitHub.')

    doc.add_heading('Awards', level=1)
    doc.add_paragraph('Employee of the Year (2022) - Tech Solutions Inc.')
    doc.add_paragraph("Dean's List (2016, 2017, 2018) - University of Technology")

    doc.save('sample_cv.docx')

if __name__ == "__main__":
    create_sample_cv()
